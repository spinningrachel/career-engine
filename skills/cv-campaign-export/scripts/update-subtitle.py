#!/usr/bin/env python3
"""
update-subtitle.py
Updates the role-specific Subtitle in a CV DOCX file's first-page header.
Usage: python update-subtitle.py <docx_path> <tagline>
"""
import sys
from docx import Document


def update_subtitle(docx_path: str, tagline: str) -> bool:
    doc = Document(docx_path)
    updated = False
    for section in doc.sections:
        for para in section.first_page_header.paragraphs:
            if para.style.name == "Subtitle":
                para.clear()
                para.add_run(tagline)
                updated = True
                break
        if updated:
            break
    if not updated:
        raise RuntimeError(
            "Subtitle style not found in first_page_header — check that the .dotx template "
            "contains a Subtitle-styled paragraph in its first-page header."
        )
    doc.save(docx_path)
    return True


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python update-subtitle.py <docx_path> <tagline>")
        sys.exit(1)
    path, tagline = sys.argv[1], sys.argv[2]
    update_subtitle(path, tagline)
    print(f"Subtitle updated: {tagline!r}")
