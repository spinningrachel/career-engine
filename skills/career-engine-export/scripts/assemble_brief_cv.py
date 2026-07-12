#!/usr/bin/env python3
"""
Post-processing assembly script for the Brief CV template (one-page, two-column).

Fills the pre-built table shell in cv-brief.dotx directly via python-docx,
bypassing pandoc's docx table writer entirely for this template -- pandoc
cannot represent the nested/merged-cell structure this layout needs (see
career-engine-export/SKILL.md, "CV -- custom-style annotation reference --
Brief variant" for the empirical findings that ruled out a pandoc-table
approach). Role-row count is dynamic: the table grows or shrinks a role row
at a time to match however many roles are passed in.

Content sources (career-engine-export/SKILL.md Step 6 documents the full
wiring):
  - skills / summary / roles  <- parsed from cv-writer's own linear markdown
    (the <!-- SIDEBAR -->-wrapped ## SKILLS block, ## PROFILE SUMMARY, and
    ## EXPERIENCE's RoleTitle/RoleActivitySingle divs). cv-writer never emits
    a table -- this is the same linear markdown format it already produces.
  - education / languages     <- parsed from $CV_FOOTER (static-cv-footer.md),
    the same file Detailed appends verbatim; Brief instead routes it into the
    sidebar cell. --cv-footer is OPTIONAL (2026-07-12 fix, cv_footer.inject in
    pipeline-preferences.json): omitted entirely from the sidebar when not
    supplied, the same "not this pipeline's job" convention as --additional
    below -- the user manages Education/Languages herself outside the pipeline.
  - name / tagline / contact  <- already-resolved pipeline/career-data values,
    passed in as CLI arguments. cv-writer never emits these (same convention
    as Detailed, whose header contact is baked into the user's own .dotx, not
    pipeline content).
  - additional                <- optional; no pipeline source resolves this
    today for any CV Type (Detailed's ## ADDITIONAL is added later by the
    user's own Word macro, post-export, per career-engine-export/SKILL.md).
    Omitted entirely from the sidebar when not supplied -- not a regression,
    a pre-existing gap.

*** SCOPE -- read before pointing this at a personalized cv-brief.dotx ***
This script hard-codes the DEFAULT template's exact table shape: a specific
row/column layout, specific merged cells (row 0 and row 2 gridSpan across
columns 1-2, column 0 vertically merged from row 1 down), and specific named
paragraph/character styles (RoleTitle, RoleActivitySingle, SkillsHeading,
Skills, PersonalDetails, ColorEmphasis, Heading 1, Subtitle, Heading 2,
Normal). It is safe to point at a personalized cv-brief.dotx ONLY when that
file is a cosmetic derivative of the default -- fonts, colors, spacing,
alignment changed, but the table's rows/columns/merges and style NAMES left
alone. It is NOT safe to point at a template where the table itself was
restructured (rows/columns added or removed, merges changed, styles renamed
or deleted) -- this script does not know how to fill a different shape, and
`validate_template_structure()` below exists specifically to fail loudly and
specifically in that case rather than silently writing into the wrong cell
or crashing on an unhelpful python-docx KeyError.
"""
import argparse
import copy
import io
import json
import subprocess
import zipfile

import docx
from lxml import etree

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
FIRST_ROLE_ROW = 3  # rows 0,1,2 are fixed: header, summary, EXPERIENCE banner
BORDER_COLOR = "B9BA00"

REQUIRED_STYLES = [
    "Heading 1", "Subtitle", "SkillsHeading", "PersonalDetails", "Skills",
    "Heading 2", "Normal", "RoleTitle", "RoleActivitySingle", "ColorEmphasis",
]


class TemplateStructureError(RuntimeError):
    """Raised when the template's table shape or styles don't match what this
    script hard-codes -- see the module docstring's SCOPE section. This is
    deliberately its own exception type (not a bare python-docx KeyError or
    IndexError) so the failure is unambiguous and actionable rather than a
    stack trace into internal cell-indexing code."""


def validate_template_structure(doc, template_path):
    """Fail loudly and specifically if this template's table doesn't match
    the shape assemble_brief_cv.py hard-codes, rather than silently writing
    into the wrong cell or crashing on an opaque python-docx exception deep in
    fill_brief_cv(). Safe to run against the default template or a purely
    cosmetic derivative (fonts/colors/alignment changed, structure and style
    names untouched) -- NOT safe to skip if the table itself was restructured."""
    problems = []

    if not doc.tables:
        raise TemplateStructureError(
            f"'{template_path}' has no tables at all. This script only supports the "
            "default cv-brief.dotx table shape (or a cosmetic-only derivative -- "
            "fonts/colors/alignment changed, table structure untouched). If this "
            "template's table was removed or restructured, this script cannot be used "
            "as-is -- see the SCOPE note at the top of assemble_brief_cv.py."
        )

    t = doc.tables[0]
    tbl = t._tbl
    trs = tbl.findall(f"{W_NS}tr")

    if len(trs) < FIRST_ROLE_ROW + 1:
        problems.append(
            f"expected at least {FIRST_ROLE_ROW + 1} rows (3 fixed rows + at least 1 role "
            f"row to use as a clone template), found {len(trs)}"
        )
    else:
        gridspan0 = trs[0].findall(f"{W_NS}tc")[1].find(f"{W_NS}tcPr/{W_NS}gridSpan")
        if gridspan0 is None or gridspan0.get(f"{W_NS}val") != "2":
            problems.append("row 0's second cell is not gridSpan=2 (expected the name/tagline header merged across columns 1-2)")

        vmerge1 = trs[1].findall(f"{W_NS}tc")[0].find(f"{W_NS}tcPr/{W_NS}vMerge")
        if vmerge1 is None:
            problems.append("row 1's first cell has no vMerge (expected the sidebar cell to start its vertical merge here)")

        gridspan2 = trs[2].findall(f"{W_NS}tc")[1].find(f"{W_NS}tcPr/{W_NS}gridSpan")
        if gridspan2 is None or gridspan2.get(f"{W_NS}val") != "2":
            problems.append("row 2's second cell is not gridSpan=2 (expected the EXPERIENCE banner merged across columns 1-2)")

        role_tcs = trs[FIRST_ROLE_ROW].findall(f"{W_NS}tc")
        if len(role_tcs) != 3:
            problems.append(f"row {FIRST_ROLE_ROW} (first role row) has {len(role_tcs)} cells, expected 3 (sidebar continuation, date, title+bullets)")
        elif role_tcs[0].find(f"{W_NS}tcPr/{W_NS}vMerge") is None:
            problems.append(f"row {FIRST_ROLE_ROW}'s first cell has no vMerge (expected the sidebar's vertical merge to continue here)")

    style_names = {s.name for s in doc.styles}
    missing_styles = [s for s in REQUIRED_STYLES if s not in style_names]
    if missing_styles:
        problems.append(f"missing required style(s): {', '.join(missing_styles)}")

    if problems:
        raise TemplateStructureError(
            f"'{template_path}' does not match the table structure/styles this script "
            "expects. This script only supports the default cv-brief.dotx (or a "
            "cosmetic-only derivative -- fonts/colors/alignment changed, table "
            "structure and style names untouched). If this template's table was "
            "restructured (rows/columns added or removed, merges changed, styles "
            "renamed or deleted), this script cannot be used as-is -- see the SCOPE "
            "note at the top of assemble_brief_cv.py. Specific problems found:\n  - "
            + "\n  - ".join(problems)
        )


# ---------------------------------------------------------------------------
# .dotx loading
# ---------------------------------------------------------------------------

def load_docx_or_dotx(path):
    """python-docx's Document() rejects a .dotx outright: it checks the main
    part's content type, not the file extension, and a .dotx's main part is
    registered as '...wordprocessingml.template.main+xml' rather than
    '...wordprocessingml.document.main+xml' -- otherwise identical
    WordprocessingML. Confirmed against the real cv-brief.dotx shipped in
    this plugin: python-docx raises ValueError ("is not a Word file") without
    this patch. Fix: rewrite the content-type override in memory (no temp
    file) before handing the bytes to python-docx."""
    with open(path, "rb") as f:
        src = f.read()
    in_buf = io.BytesIO(src)
    out_buf = io.BytesIO()
    with zipfile.ZipFile(in_buf, "r") as zin, zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "[Content_Types].xml":
                data = data.replace(
                    b"wordprocessingml.template.main+xml",
                    b"wordprocessingml.document.main+xml",
                )
            zout.writestr(item, data)
    out_buf.seek(0)
    return docx.Document(out_buf)


# ---------------------------------------------------------------------------
# Cell / run writing
# ---------------------------------------------------------------------------

def _apply_run_style(run, style):
    """style is None (plain run, paragraph style governs it), 'bold' (direct
    character formatting, no named style), or the name of a character style
    defined in the template (e.g. ColorEmphasis)."""
    if style is None:
        return
    if style == "bold":
        run.bold = True
    else:
        run.style = style


def _clear_paragraph_content(p):
    """Remove every child of the paragraph except pPr (paragraph properties).
    p.runs alone is not enough: python-docx's `paragraph.runs` only sees
    direct <w:r> children and misses runs wrapped in <w:hyperlink> (or a
    bookmark) -- confirmed against this exact template, whose sidebar
    'www.example.com' placeholder line is hyperlink-wrapped. Clearing only
    p.runs left that text behind, silently concatenated onto whatever this
    function wrote next into the reused paragraph."""
    pPr_qn = f"{W_NS}pPr"
    for child in list(p._p):
        if child.tag != pPr_qn:
            p._p.remove(child)


def set_cell_paragraphs(cell, entries):
    """entries: list of (content, para_style) tuples. `content` is either a
    plain string (single unstyled run) or a list of (run_text, run_style)
    tuples for a paragraph that mixes styles -- e.g. a RoleTitle line with
    one ColorEmphasis-styled company name, or the Earlier line with a bold
    "Earlier:" prefix. Replaces all paragraphs in the cell; callers must pass
    the complete desired content."""
    paras = cell.paragraphs
    for i, (content, style) in enumerate(entries):
        p = paras[i] if i < len(paras) else cell.add_paragraph()
        _clear_paragraph_content(p)
        runs = content if isinstance(content, list) else [(content, None)]
        for text, run_style in runs:
            run = p.add_run(text)
            _apply_run_style(run, run_style)
        p.style = style
    for extra in paras[len(entries):]:
        extra._element.getparent().remove(extra._element)


def _set_col0_bottom_border(tr, on: bool):
    """Add/remove the closing bottom border on a role row's sidebar (col0) cell."""
    tc0 = tr.findall(f"{W_NS}tc")[0]
    tcPr = tc0.find(f"{W_NS}tcPr")
    borders = tcPr.find(f"{W_NS}tcBorders")
    if on:
        if borders is None:
            borders = etree.SubElement(tcPr, f"{W_NS}tcBorders")
        bottom = borders.find(f"{W_NS}bottom")
        if bottom is None:
            bottom = etree.SubElement(borders, f"{W_NS}bottom")
        bottom.set(f"{W_NS}val", "single")
        bottom.set(f"{W_NS}sz", "6")
        bottom.set(f"{W_NS}space", "0")
        bottom.set(f"{W_NS}color", BORDER_COLOR)
    else:
        if borders is not None:
            bottom = borders.find(f"{W_NS}bottom")
            if bottom is not None:
                borders.remove(bottom)
            if len(borders) == 0:
                tcPr.remove(borders)


def set_role_row_count(table, n: int):
    """Grow or shrink the table so there are exactly n role rows (rows 3..3+n-1)."""
    tbl = table._tbl
    all_trs = tbl.findall(f"{W_NS}tr")
    current_n = len(all_trs) - FIRST_ROLE_ROW
    if n == current_n:
        pass
    elif n > current_n:
        # clone a middle role row (index 4 if it exists, else the last role row) --
        # anything without a closing bottom border on col0 -- and append the copies.
        template_row = all_trs[min(FIRST_ROLE_ROW + 1, len(all_trs) - 1)]
        for _ in range(n - current_n):
            new_tr = copy.deepcopy(template_row)
            _set_col0_bottom_border(new_tr, on=False)
            tbl.append(new_tr)
    else:
        for tr in all_trs[FIRST_ROLE_ROW + n:]:
            tbl.remove(tr)

    # Fix up the closing border: only the true last row's col0 gets the bottom border.
    all_trs = tbl.findall(f"{W_NS}tr")
    for i, tr in enumerate(all_trs[FIRST_ROLE_ROW:]):
        is_last = (i == n - 1)
        _set_col0_bottom_border(tr, on=is_last)


# ---------------------------------------------------------------------------
# Markdown -> data parsing (via pandoc's own JSON AST, not regex/manual typing --
# the AST is immune to the line-wrapping corruption that broke the annotation-only
# approaches during testing, since we parse cv-writer's *source* markdown, never
# a pandoc-rendered/re-wrapped round trip).
# ---------------------------------------------------------------------------

def _pandoc_ast(md_text):
    result = subprocess.run(
        ["pandoc", "-f", "markdown", "-t", "json"],
        input=md_text, capture_output=True, text=True, encoding="utf-8", check=True,
    )
    return json.loads(result.stdout)


def _node_text(node):
    t = node.get("t")
    if t == "Str":
        return node["c"]
    if t in ("Space", "SoftBreak"):
        return " "
    if t == "Span":
        return _inline_text(node["c"][1])
    if t in ("Strong", "Emph"):
        return _inline_text(node["c"])
    return ""


def _inline_text(inlines):
    return "".join(_node_text(n) for n in inlines).strip()


def _custom_style_of(span_node):
    for k, v in span_node["c"][0][2]:
        if k == "custom-style":
            return v
    return None


def _inline_runs(inlines):
    """Flatten inline content into (text, run_style) runs. A Span carrying a
    custom-style attribute (e.g. ColorEmphasis on a company name) becomes its
    own tagged run; a Strong node (e.g. the "Earlier:" prefix) becomes its own
    run tagged 'bold'; everything else flattens to plain text, merged with
    adjacent plain text into a single unstyled run."""
    runs = []
    buf = []

    def flush():
        if buf:
            runs.append(("".join(buf), None))
            buf.clear()

    for node in inlines:
        t = node.get("t")
        if t == "Span":
            flush()
            runs.append((_inline_text(node["c"][1]), _custom_style_of(node)))
        elif t == "Strong":
            flush()
            runs.append((_inline_text(node["c"]), "bold"))
        elif t == "Str":
            buf.append(node["c"])
        elif t in ("Space", "SoftBreak"):
            buf.append(" ")
        elif t == "Emph":
            buf.append(_inline_text(node["c"]))
        # anything else is not part of this controlled content vocabulary; ignored.
    flush()
    return runs


def _div_style(block):
    return dict(block["c"][0][2]).get("custom-style")


def parse_cv_markdown(md_text):
    """Parse cv-writer's Brief-CV linear markdown into {skills, summary, roles}.
    cv-writer emits exactly: a <!-- SIDEBAR --> ... <!-- /SIDEBAR --> wrapped
    ## SKILLS block, a ## PROFILE SUMMARY paragraph, and a ## EXPERIENCE
    section of RoleTitle/RoleActivitySingle divs closed by an optional
    **Earlier:** paragraph (career-engine-export/SKILL.md Brief annotation
    reference; agents/cv-writer.md Brief-Specific Rules). Never a table --
    that path was proven unreliable through pandoc's docx writer, which is
    exactly why this script exists instead of a pandoc conversion.

    roles: [{"title": run-list|None, "dates": str, "bullets": [str|run-list, ...]}]
    "title" is None only for the closing Earlier row, whose one bullet entry
    is itself a run-list (bold "Earlier:" prefix + plain rest).
    """
    ast = _pandoc_ast(md_text)
    blocks = ast["blocks"]

    skills = None
    summary_parts = []
    roles = []
    in_sidebar = False
    section = None  # None | "summary" | "experience"
    current_role = None

    for block in blocks:
        t = block.get("t")

        if t == "RawBlock" and block["c"][0] == "html":
            html = block["c"][1].strip()
            if html == "<!-- SIDEBAR -->":
                in_sidebar = True
            elif html == "<!-- /SIDEBAR -->":
                in_sidebar = False
            continue

        if t == "Header" and block["c"][0] == 2:
            heading = _inline_text(block["c"][2]).upper()
            section = heading if heading in ("PROFILE SUMMARY", "EXPERIENCE") else None
            continue

        if t == "Div":
            style = _div_style(block)
            content_blocks = block["c"][1]
            first_para_inlines = content_blocks[0]["c"] if content_blocks else []

            if in_sidebar and style == "Skills":
                skills = _inline_text(first_para_inlines)
                continue

            if section == "EXPERIENCE" and style == "RoleTitle":
                if current_role is not None:
                    roles.append(current_role)
                current_role = {"title": _inline_runs(first_para_inlines), "dates": None, "bullets": []}
                continue

            if section == "EXPERIENCE" and style == "RoleActivitySingle" and current_role is not None:
                text = _inline_text(first_para_inlines)
                if current_role["dates"] is None:
                    current_role["dates"] = text
                else:
                    current_role["bullets"].append(text)
                continue

            continue

        if t == "Para" and section == "PROFILE SUMMARY":
            summary_parts.append(_inline_text(block["c"]))
            continue

        if t == "Para" and section == "EXPERIENCE":
            # The one top-level (non-Div) paragraph cv-writer ever emits inside
            # EXPERIENCE is the closing **Earlier:** line (writer-craft/SKILL.md §5b).
            if current_role is not None:
                roles.append(current_role)
                current_role = None
            roles.append({"title": None, "dates": "", "bullets": [_inline_runs(block["c"])]})
            continue

    if current_role is not None:
        roles.append(current_role)

    if not roles:
        raise RuntimeError(
            "parse_cv_markdown found zero roles -- check that the markdown has a "
            "'## EXPERIENCE' heading (level-2) with RoleTitle/RoleActivitySingle "
            "custom-style divs beneath it, per career-engine-export/SKILL.md's "
            "Brief annotation reference."
        )

    return {
        "skills": skills or "",
        "summary": "\n\n".join(summary_parts),
        "roles": roles,
    }


def parse_footer_markdown(footer_text):
    """Parse $CV_FOOTER (static-cv-footer.md) into {education: [...], languages: str}.
    Same source file as Detailed -- for Brief, the sidebar routes this content
    instead of appending it at the end of the document (career-engine-export/SKILL.md,
    Templates section). Footer convention: '## EDUCATION' with one paragraph per
    degree -- kept one-per-line, matching the template's own per-degree 'Skills'-
    styled paragraphs. '## LANGUAGES' with one paragraph per language -- these are
    joined into a single pipe-separated line, because the Brief template's sidebar
    renders Languages as ONE flat 'Skills'-styled paragraph (confirmed against the
    real cv-brief.dotx default content), unlike the footer file's own per-line
    convention."""
    ast = _pandoc_ast(footer_text)
    blocks = ast["blocks"]

    education = []
    language_lines = []
    section = None

    for block in blocks:
        t = block.get("t")
        if t == "Header" and block["c"][0] == 2:
            heading = _inline_text(block["c"][2]).upper()
            section = heading if heading in ("EDUCATION", "LANGUAGES") else None
            continue
        if t == "Para":
            if section == "EDUCATION":
                education.append(_inline_runs(block["c"]))
            elif section == "LANGUAGES":
                language_lines.append(_inline_text(block["c"]))

    return {
        "education": education,
        "languages": " | ".join(language_lines),
    }


# ---------------------------------------------------------------------------
# Assembly
# ---------------------------------------------------------------------------

def fill_brief_cv(template_path, output_path, data):
    doc = load_docx_or_dotx(template_path)
    validate_template_structure(doc, template_path)
    t = doc.tables[0]

    set_role_row_count(t, len(data["roles"]))

    set_cell_paragraphs(t.cell(0, 1), [
        (data["name"], "Heading 1"),
        (data["tagline"], "Subtitle"),
    ])

    sidebar_entries = [("CONTACT details", "SkillsHeading")]
    for line in data["contact"]:
        sidebar_entries.append((line, "PersonalDetails"))
    sidebar_entries.append(("SKILLS", "SkillsHeading"))
    sidebar_entries.append((data["skills"], "Skills"))
    if data.get("languages"):
        sidebar_entries.append(("LANGUAGES", "SkillsHeading"))
        sidebar_entries.append((data["languages"], "Skills"))
    if data.get("additional"):
        sidebar_entries.append(("ADDITIONAL", "SkillsHeading"))
        sidebar_entries.append((data["additional"], "Skills"))
    if data.get("education"):
        sidebar_entries.append(("EDUCATION", "Heading 2"))
        for entry in data["education"]:
            sidebar_entries.append((entry, "Skills"))
    set_cell_paragraphs(t.cell(1, 0), sidebar_entries)

    set_cell_paragraphs(t.cell(1, 1), [(data["summary"], "Normal")])
    set_cell_paragraphs(t.cell(2, 1), [("professional EXPERIENCE", "Heading 2")])

    for i, role in enumerate(data["roles"]):
        r = FIRST_ROLE_ROW + i
        set_cell_paragraphs(t.cell(r, 1), [(role["dates"], "RoleActivitySingle")])
        role_entries = []
        if role["title"] is not None:
            role_entries.append((role["title"], "RoleTitle"))
        for bullet in role["bullets"]:
            role_entries.append((bullet, "RoleActivitySingle"))
        set_cell_paragraphs(t.cell(r, 2), role_entries)

    doc.save(output_path)


def build_data(args, cv_parsed, footer_parsed):
    return {
        "name": args.name,
        "tagline": args.tagline,
        "contact": args.contact,
        "skills": cv_parsed["skills"],
        "languages": footer_parsed["languages"],
        "additional": args.additional,
        "education": footer_parsed["education"],
        "summary": cv_parsed["summary"],
        "roles": cv_parsed["roles"],
    }


def main():
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--cv-md", required=True, help="Path to cv-writer's Brief CV markdown")
    ap.add_argument("--cv-footer", default=None, help="Path to $CV_FOOTER (static-cv-footer.md); omit if cv_footer.inject is false in pipeline-preferences.json")
    ap.add_argument("--template", required=True, help="Path to $CV_TEMPLATE_BRIEF (cv-brief.dotx)")
    ap.add_argument("--output", required=True, help="Output .docx path")
    ap.add_argument("--name", required=True, help="User's full name (career-data identity placeholder, already resolved)")
    ap.add_argument("--tagline", required=True, help="Role title, verbatim from JD -- same value update-subtitle.py uses for Detailed")
    ap.add_argument("--contact", action="append", default=[], help="One contact line (repeatable) -- e.g. city/country, phone, email, site, LinkedIn")
    ap.add_argument("--additional", default=None, help="Optional pipe-separated Additional line; omitted entirely if not passed (no established content source yet -- see career-engine-export/SKILL.md)")
    args = ap.parse_args()

    with open(args.cv_md, encoding="utf-8") as f:
        cv_parsed = parse_cv_markdown(f.read())
    if args.cv_footer:
        with open(args.cv_footer, encoding="utf-8") as f:
            footer_parsed = parse_footer_markdown(f.read())
    else:
        footer_parsed = {"education": [], "languages": ""}

    data = build_data(args, cv_parsed, footer_parsed)
    fill_brief_cv(args.template, args.output, data)
    print(f"Brief CV assembled: {args.output} ({len(data['roles'])} role rows)")


if __name__ == "__main__":
    main()
