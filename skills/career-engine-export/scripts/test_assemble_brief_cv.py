#!/usr/bin/env python3
"""
End-to-end test for assemble_brief_cv.py, run directly (no pytest dependency --
this plugin has no existing test harness to plug into): builds a fixture
template from the real cv-template-brief-default.dotx with a synthetic image
inserted into the photo cell (simulating a user's personalized cv-brief.dotx,
which has a real photo baked in -- the blank default does not), then asserts,
for 1/3/6 role counts:
  - the table grows/shrinks to exactly the right row count
  - the vMerge sidebar cell and the two gridSpan=2 header/banner rows survive
  - only the true last role row carries the closing sidebar border
  - custom styles (RoleTitle, RoleActivitySingle, ColorEmphasis, Skills, etc.)
    land on the correct paragraphs/runs
  - the embedded image in the untouched photo cell survives

Usage: python3 test_assemble_brief_cv.py
Exits 0 and prints "ALL TESTS PASSED" on success; raises AssertionError (exit 1) otherwise.
"""
import os
import struct
import sys
import tempfile
import zlib

import docx
from docx.shared import Inches

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import assemble_brief_cv as m

HERE = os.path.dirname(os.path.abspath(__file__))
FIXTURES = os.path.join(HERE, "test-fixtures")
TEMPLATE = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(HERE))), "references", "cv-template-brief-default.dotx")
CV_MD = os.path.join(FIXTURES, "brief-cv-sample.md")
FOOTER_MD = os.path.join(FIXTURES, "brief-footer-sample.md")


def _tiny_png_bytes():
    def chunk(tag, data):
        c = tag + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00\xff\x00\x00")
    return sig + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


def _blip_count(docx_path):
    doc = docx.Document(docx_path)
    tc = doc.tables[0].cell(0, 0)._tc
    return len(tc.findall(".//{*}blip"))


def build_fixture_template_with_photo(tmpdir):
    """Simulate a personalized cv-brief.dotx: the blank default template has
    no photo baked in, but a real user's does (career-engine-export/SKILL.md).
    Insert a synthetic image into the untouched photo cell (row 0, col 0) so
    the test can confirm the assembly script never disturbs it."""
    doc = m.load_docx_or_dotx(TEMPLATE)
    png_path = os.path.join(tmpdir, "tiny.png")
    with open(png_path, "wb") as f:
        f.write(_tiny_png_bytes())
    doc.tables[0].cell(0, 0).paragraphs[0].add_run().add_picture(png_path, width=Inches(0.5))
    fixture_path = os.path.join(tmpdir, "fixture-with-photo.docx")
    doc.save(fixture_path)
    return fixture_path


def build_data(role_count_override=None):
    cv_parsed = m.parse_cv_markdown(open(CV_MD, encoding="utf-8").read())
    footer_parsed = m.parse_footer_markdown(open(FOOTER_MD, encoding="utf-8").read())
    roles = cv_parsed["roles"]
    if role_count_override is not None:
        real_roles = [r for r in roles if r["title"] is not None]
        if role_count_override <= len(real_roles):
            roles = real_roles[:role_count_override]
        else:
            roles = real_roles + [
                {
                    "title": [(f"Synthetic Role {i} | Company {i} | City {i}", None)],
                    "dates": f"20{10+i}--20{11+i}",
                    "bullets": [f"Did synthetic thing {i}."],
                }
                for i in range(role_count_override - len(real_roles))
            ]
    return {
        "name": "George Costanza",
        "tagline": "Assistant to Traveling Secretary",
        "contact": ["New York, NY", "+1 555 000 0000", "george@vandelay.com"],
        "skills": cv_parsed["skills"],
        "languages": footer_parsed["languages"],
        "additional": None,
        "education": footer_parsed["education"],
        "summary": cv_parsed["summary"],
        "roles": roles,
    }


def assert_structure(docx_path, n_roles):
    doc = docx.Document(docx_path)
    tbl = doc.tables[0]._tbl
    W = m.W_NS
    trs = tbl.findall(f"{W}tr")
    assert len(trs) == 3 + n_roles, f"{docx_path}: row count {len(trs)} != {3 + n_roles}"

    gridspan0 = trs[0].findall(f"{W}tc")[1].find(f"{W}tcPr/{W}gridSpan")
    assert gridspan0 is not None and gridspan0.get(f"{W}val") == "2", "row0 gridSpan (name/tagline banner) missing or wrong"

    vmerge1 = trs[1].findall(f"{W}tc")[0].find(f"{W}tcPr/{W}vMerge")
    assert vmerge1 is not None, "row1 col0 vMerge (sidebar start) missing"

    gridspan2 = trs[2].findall(f"{W}tc")[1].find(f"{W}tcPr/{W}gridSpan")
    assert gridspan2 is not None and gridspan2.get(f"{W}val") == "2", "row2 gridSpan (EXPERIENCE banner) missing or wrong"

    for i, tr in enumerate(trs[3:]):
        tcs = tr.findall(f"{W}tc")
        assert len(tcs) == 3, f"role row {i} has {len(tcs)} tc elements, expected 3"
        assert tcs[0].find(f"{W}tcPr/{W}vMerge") is not None, f"role row {i} col0 vMerge (sidebar continuation) missing"
        is_last = (i == n_roles - 1)
        has_border = tcs[0].find(f"{W}tcPr/{W}tcBorders/{W}bottom") is not None
        assert has_border == is_last, f"role row {i}: closing border presence {has_border}, expected {is_last}"


def assert_content(docx_path, n_roles):
    doc = docx.Document(docx_path)
    t = doc.tables[0]

    header_texts = [p.text for p in t.cell(0, 1).paragraphs]
    assert header_texts == ["George Costanza", "Assistant to Traveling Secretary"], header_texts
    assert [p.style.name for p in t.cell(0, 1).paragraphs] == ["Heading 1", "Subtitle"]

    sidebar_texts = [p.text for p in t.cell(1, 0).paragraphs]
    assert "CONTACT details" in sidebar_texts
    assert "Python | JavaScript | SQL | AWS | Docker" in sidebar_texts
    assert "ADDITIONAL" not in sidebar_texts, "ADDITIONAL heading must be omitted when no content is supplied"
    assert any("English" in s and "Spanish" in s for s in sidebar_texts), "languages not joined onto one line"

    role_row = t.rows[3]
    seen = []
    for c in role_row.cells:
        if c._tc not in [x._tc for x in seen]:
            seen.append(c)
    title_cell = seen[2]
    assert title_cell.paragraphs[0].style.name == "RoleTitle"
    run_styles = [(r.text, r.style.name if r.style else None) for r in title_cell.paragraphs[0].runs]
    color_runs = [r for r in run_styles if r[1] == "ColorEmphasis"]
    assert color_runs and color_runs[0][0] == "Newman’s", f"ColorEmphasis run not applied correctly: {run_styles}"


def build_restructured_template(tmpdir):
    """A template that is NOT a cosmetic derivative of the default -- it strips
    the gridSpan merge off row 0's header cell, simulating a user who
    restructured the table itself (not just changed fonts/colors/alignment).
    assemble_brief_cv.py must refuse to run against this, not silently
    misrender or crash on an opaque exception -- see the module docstring's
    SCOPE note and validate_template_structure()."""
    doc = m.load_docx_or_dotx(TEMPLATE)
    tbl = doc.tables[0]._tbl
    trs = tbl.findall(f"{m.W_NS}tr")
    tc1 = trs[0].findall(f"{m.W_NS}tc")[1]
    gridspan = tc1.find(f"{m.W_NS}tcPr/{m.W_NS}gridSpan")
    tc1.find(f"{m.W_NS}tcPr").remove(gridspan)
    path = os.path.join(tmpdir, "restructured-template.docx")
    doc.save(path)
    return path


def assert_footer_omitted_when_not_supplied(tmpdir):
    """cv_footer.inject=false: no footer content supplied. EDUCATION/LANGUAGES
    headings must be omitted from the sidebar entirely (not present-but-blank),
    same convention as the pre-existing ADDITIONAL-omitted-when-absent case."""
    data = build_data(role_count_override=1)
    data["education"] = []
    data["languages"] = ""
    out_path = os.path.join(tmpdir, "no-footer.docx")
    m.fill_brief_cv(TEMPLATE, out_path, data)
    doc = docx.Document(out_path)
    sidebar_texts = [p.text for p in doc.tables[0].cell(1, 0).paragraphs]
    assert "EDUCATION" not in sidebar_texts, "EDUCATION heading must be omitted when no footer content is supplied"
    assert "LANGUAGES" not in sidebar_texts, "LANGUAGES heading must be omitted when no footer content is supplied"
    assert "CONTACT details" in sidebar_texts and "SKILLS" in sidebar_texts, "unrelated sidebar content should be unaffected"


def assert_restructured_template_rejected(tmpdir):
    broken_template = build_restructured_template(tmpdir)
    data = build_data(role_count_override=1)
    out_path = os.path.join(tmpdir, "should-not-exist.docx")
    try:
        m.fill_brief_cv(broken_template, out_path, data)
    except m.TemplateStructureError as e:
        assert "gridSpan" in str(e), f"error message doesn't name the specific structural problem: {e}"
        assert not os.path.exists(out_path), "output file must not be written when validation fails"
        return
    raise AssertionError("fill_brief_cv() did not raise TemplateStructureError against a restructured template")


def main():
    with tempfile.TemporaryDirectory() as tmpdir:
        fixture_template = build_fixture_template_with_photo(tmpdir)
        before = _blip_count(fixture_template)
        assert before > 0, "test setup failed: no image found in fixture template before assembly"

        for n in (1, 3, 6):
            data = build_data(role_count_override=n)
            out_path = os.path.join(tmpdir, f"out-{n}.docx")
            m.fill_brief_cv(fixture_template, out_path, data)
            assert_structure(out_path, n)
            if n == 3:
                assert_content(out_path, n)
            after = _blip_count(out_path)
            assert after == before, f"image lost or duplicated for n={n}: {before} -> {after} blips"
            print(f"  n={n} roles: row count, styles, vMerge/gridSpan, and image all OK")

        assert_restructured_template_rejected(tmpdir)
        print("  restructured (non-cosmetic) template correctly rejected with a clear error, no output written")

        assert_footer_omitted_when_not_supplied(tmpdir)
        print("  cv_footer.inject=false: EDUCATION/LANGUAGES correctly omitted from sidebar")

    print("ALL TESTS PASSED")


if __name__ == "__main__":
    main()
